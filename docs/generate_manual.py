#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成分发版 Word 操作手册"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_PATH = os.path.join(BASE_DIR, '项目信息一键提取_操作手册.docx')

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_image(doc, img_path, caption=None, width=6.0):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
    else:
        add_para(doc, f'[截图占位：{caption}]', italic=True, color=RGBColor(255, 0, 0))

def main():
    doc = Document()

    # 标题
    title = doc.add_heading('项目信息一键提取 — 操作手册', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, '本文档面向最终用户，介绍如何安装、配置和使用本工具。', italic=True)
    add_para(doc, '注意：截图中的路径、密码等仅为示例，实际使用请替换为本地真实配置。', italic=True, color=RGBColor(255, 0, 0))

    # 一、环境要求
    add_heading(doc, '一、环境要求', level=1)
    add_para(doc, 'Windows 10/11')
    add_para(doc, 'Node.js 18+（建议 20+）')
    add_para(doc, 'MySQL 5.7+ 或 MariaDB 10+')
    add_para(doc, 'Chrome 浏览器（支持 Manifest V3）')

    # 二、目录说明
    add_heading(doc, '二、目录说明', level=1)
    add_para(doc, '解压分发包后，主要目录结构如下：')
    add_para(doc, 'backend/：Node.js 后端服务', bold=True)
    add_para(doc, 'chrome-extension/：Chrome 插件', bold=True)
    add_para(doc, '使用说明.md / 本手册：操作文档', bold=True)
    add_para(doc, '首次使用前，需要进入 backend 目录执行 npm install 安装依赖。')

    # 三、启动后端
    add_heading(doc, '三、启动后端服务', level=1)
    add_para(doc, '1. 打开命令行，进入 backend 目录：')
    add_para(doc, 'cd backend', bold=True)
    add_para(doc, '2. 安装依赖（首次）：')
    add_para(doc, 'npm install', bold=True)
    add_para(doc, '3. 启动服务：')
    add_para(doc, 'npm start', bold=True)
    add_para(doc, '4. 浏览器访问 http://localhost:3000/ 查看是否正常运行。')

    # 四、首次配置
    add_heading(doc, '四、首次配置', level=1)
    add_para(doc, '点击查询页面右上角「设置」，进入设置页面：')
    add_image(doc, os.path.join(BASE_DIR, 'screenshots', 'query_page.png'),
              '图 1：项目信息查询页面（右上角为「设置」入口）', width=6.2)

    add_para(doc, '在设置页面填写：')
    add_para(doc, '本地归档：归档文件夹、浏览器默认下载目录')
    add_para(doc, '数据库：主机、端口、用户名、密码、数据库名')
    add_para(doc, '填写后先点击「测试连接」，再点击「保存设置」。')
    add_image(doc, os.path.join(BASE_DIR, 'screenshots', 'settings_page.png'),
              '图 2：设置页面（本地归档与数据库配置）', width=6.2)

    # 五、安装 Chrome 插件
    add_heading(doc, '五、安装 Chrome 插件', level=1)
    add_para(doc, '1. 打开 Chrome，地址栏输入 chrome://extensions/')
    add_para(doc, '2. 右上角开启「开发者模式」')
    add_para(doc, '3. 点击「加载已解压的扩展程序」')
    add_para(doc, '4. 选择分发包中的 chrome-extension 文件夹')
    add_para(doc, '5. 安装成功后，浏览器右上角会出现插件图标')
    add_image(doc, os.path.join(BASE_DIR, 'screenshots', 'placeholder_extension.png'),
              '图 3：Chrome 扩展管理页（请在此处替换为实际截图）', width=5.0)

    # 六、日常使用
    add_heading(doc, '六、日常使用流程', level=1)
    add_para(doc, '1. 确保后端服务已启动。')
    add_para(doc, '2. 用 Chrome 打开 CPMS 立项批复页面。')
    add_para(doc, '3. 点击插件图标，打开侧边栏。')
    add_para(doc, '4. 插件会自动提取页面正文并填充表单；核对字段，必要时手工修改。')
    add_image(doc, os.path.join(BASE_DIR, 'screenshots', 'placeholder_sidepanel.png'),
              '图 4：插件侧边栏（请在此处替换为实际截图）', width=5.0)
    add_para(doc, '5. 点击「提交到数据库」。')
    add_para(doc, '6. 提交成功后侧边栏会自动关闭，后台自动完成：')
    add_para(doc, '   - 点击页面「表单下载」按钮')
    add_para(doc, '   - 移动并重命名下载文件到归档文件夹')
    add_para(doc, '   - 生成项目字段 Excel')

    # 七、查询与导出
    add_heading(doc, '七、查询已提交项目', level=1)
    add_para(doc, '打开 http://localhost:3000/：')
    add_para(doc, '点击表格中的「项目编码」可打开对应本地文件夹。')
    add_para(doc, '勾选记录后点击「导出 Excel」可导出选中项目。')

    # 八、常见问题
    add_heading(doc, '八、常见问题', level=1)
    add_para(doc, 'Q：后端启动报端口占用？')
    add_para(doc, 'A：结束占用 3000 端口的进程，或修改 .env 中的 PORT。')
    add_para(doc, 'Q：归档文件没有移动？')
    add_para(doc, 'A：检查设置中的归档文件夹和浏览器默认下载目录是否正确。')
    add_para(doc, 'Q：数据库连接测试失败？')
    add_para(doc, 'A：确认 MySQL 服务已启动、用户名密码正确、且用户有 project_tracker 数据库权限。')

    doc.save(DOC_PATH)
    print(f'已生成：{DOC_PATH}')

if __name__ == '__main__':
    main()
